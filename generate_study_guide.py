#!/usr/bin/env python3
"""Generate a 100-question DOCX study guide from the approved test bank."""

import json
import os
from pathlib import Path

from lxml import etree

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BANK_DIR = Path(__file__).parent / "bank" / "approved"
OUTPUT_PATH = Path(__file__).parent / "exports" / "study_guide.docx"

# Reporting-category display order
CATEGORY_ORDER = [
    "Civics and Government",
    "Geography and Economics",
    "History",
]


def load_questions():
    """Load all approved questions, sorted by category then ID."""
    questions = []
    for fp in sorted(BANK_DIR.glob("*.json")):
        with open(fp) as f:
            questions.append(json.load(f))

    # Sort: category order first, then by question id
    cat_rank = {c: i for i, c in enumerate(CATEGORY_ORDER)}
    questions.sort(key=lambda q: (cat_rank.get(q["reporting_category"], 99), q["id"]))
    return questions


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = etree.SubElement(
            tcPr, qn("w:tcBorders")
        )
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = etree.SubElement(
                tcBorders, qn(f"w:{edge}")
            )
        element.set(qn("w:val"), val.get("val", "single"))
        element.set(qn("w:sz"), val.get("sz", "4"))
        element.set(qn("w:color"), val.get("color", "000000"))
        element.set(qn("w:space"), val.get("space", "0"))


def add_header_block(doc):
    """Add title + Name / Date / Class header for paper use."""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("5th Grade Social Studies")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Study Guide  \u2014  100 Questions")
    run.bold = True
    run.font.size = Pt(14)

    # Name / Date / Class table
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels = ["Name:", "", "Date:", "", "Class:", ""]
    widths = [Inches(0.7), Inches(2.5), Inches(0.6), Inches(1.5), Inches(0.6), Inches(1.5)]

    row = table.rows[0]
    for i, (label, width) in enumerate(zip(labels, widths)):
        cell = row.cells[i]
        cell.width = width
        p = cell.paragraphs[0]
        if label:
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(11)
        # Add bottom border to blank cells for a fill-in line
        if not label:
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "6", "color": "000000", "space": "0"},
            )

    # Spacing after header
    spacer = doc.add_paragraph()
    spacer.space_after = Pt(4)


def add_category_heading(doc, category):
    """Add a bold section heading for a reporting category."""
    heading = doc.add_paragraph()
    heading.space_before = Pt(14)
    heading.space_after = Pt(4)
    run = heading.add_run(category)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # Add a thin line under heading
    border_p = heading._p
    pPr = border_p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "444444")


def add_question(doc, number, q):
    """Add a single question block to the document."""
    # Stimulus (reading passage), if present
    stimulus = q.get("stimulus")
    if stimulus and stimulus.get("content"):
        stim_para = doc.add_paragraph()
        stim_para.space_before = Pt(6)
        stim_para.space_after = Pt(2)
        stim_para.paragraph_format.left_indent = Inches(0.3)
        run = stim_para.add_run(stimulus["content"])
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Question stem
    stem_para = doc.add_paragraph()
    stem_para.space_before = Pt(4)
    stem_para.space_after = Pt(2)
    num_run = stem_para.add_run(f"{number}. ")
    num_run.bold = True
    num_run.font.size = Pt(11)
    stem_run = stem_para.add_run(q["stem"])
    stem_run.font.size = Pt(11)

    # Answer choices
    for choice in q["choices"]:
        choice_para = doc.add_paragraph()
        choice_para.space_before = Pt(0)
        choice_para.space_after = Pt(1)
        choice_para.paragraph_format.left_indent = Inches(0.5)
        run = choice_para.add_run(f"{choice['id']}.  {choice['text']}")
        run.font.size = Pt(11)

    # Small spacer between questions
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(0)
    spacer.space_after = Pt(4)


def add_answer_key(doc, questions):
    """Add an answer key section at the end."""
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Answer Key")
    run.bold = True
    run.font.size = Pt(16)
    heading.space_after = Pt(12)

    # Build answer key in a compact table (10 rows x 10 columns of Q#-Answer)
    cols = 5
    rows_needed = (len(questions) + cols - 1) // cols
    table = doc.add_table(rows=rows_needed + 1, cols=cols * 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for c in range(cols):
        cell_q = header_row.cells[c * 2]
        cell_a = header_row.cells[c * 2 + 1]
        rq = cell_q.paragraphs[0].add_run("#")
        rq.bold = True
        rq.font.size = Pt(9)
        ra = cell_a.paragraphs[0].add_run("Ans")
        ra.bold = True
        ra.font.size = Pt(9)

    # Fill answers
    for idx, q in enumerate(questions):
        correct = next(ch["id"] for ch in q["choices"] if ch["correct"])
        col_group = idx // rows_needed
        row_in_group = idx % rows_needed
        if col_group >= cols:
            break
        cell_q = table.rows[row_in_group + 1].cells[col_group * 2]
        cell_a = table.rows[row_in_group + 1].cells[col_group * 2 + 1]
        rq = cell_q.paragraphs[0].add_run(str(idx + 1))
        rq.font.size = Pt(9)
        ra = cell_a.paragraphs[0].add_run(correct)
        ra.font.size = Pt(9)


def main():
    questions = load_questions()
    print(f"Loaded {len(questions)} questions")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    add_header_block(doc)

    current_category = None
    number = 1

    for q in questions:
        cat = q["reporting_category"]
        if cat != current_category:
            add_category_heading(doc, cat)
            current_category = cat

        add_question(doc, number, q)
        number += 1

    add_answer_key(doc, questions)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Study guide saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
