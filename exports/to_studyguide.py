#!/usr/bin/env python3
"""
Export approved questions to a printable DOCX study guide.
Grouped by reporting category, essentials marked with a star,
answer key at the end.

Usage:
    python exports/to_studyguide.py
    → writes exports/studyguide.docx
"""

import json
import os
import glob as globmod

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
APPROVED_DIR = os.path.join(ROOT, "bank", "approved")
OUTPUT_PATH = os.path.join(ROOT, "exports", "studyguide.docx")

CATEGORY_ORDER = [
    "Civics and Government",
    "History",
    "Geography and Economics",
]


def load_approved():
    questions = []
    for path in sorted(globmod.glob(os.path.join(APPROVED_DIR, "*.json"))):
        with open(path, "r") as f:
            questions.append(json.load(f))
    return questions


def group_by_category(questions):
    groups = {cat: [] for cat in CATEGORY_ORDER}
    for q in questions:
        rc = q.get("reporting_category", "Unknown")
        if rc in groups:
            groups[rc].append(q)
        else:
            groups.setdefault(rc, []).append(q)
    return groups


def export():
    questions = load_approved()
    if not questions:
        print("No approved questions to export.")
        return

    doc = Document()

    # --- Styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading("ILEARN Social Studies — Practice Study Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Grade 5 | Indiana Academic Standards 2023")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacer

    groups = group_by_category(questions)
    question_num = 0
    answer_key = []

    for category in CATEGORY_ORDER:
        cat_questions = groups.get(category, [])
        if not cat_questions:
            continue

        doc.add_heading(category, level=1)

        for q in cat_questions:
            question_num += 1

            # Essential star
            essential_mark = " ★" if q.get("standard_essential") else ""

            # Standard code + question number
            header = doc.add_paragraph()
            run = header.add_run(
                f"Question {question_num}  [{q['standard_code']}{essential_mark}]"
            )
            run.bold = True
            run.font.size = Pt(11)

            # Stimulus
            if q.get("stimulus", {}).get("content"):
                stim_para = doc.add_paragraph()
                stim_run = stim_para.add_run(q["stimulus"]["content"])
                stim_run.italic = True
                stim_run.font.size = Pt(10)

            # Stem
            doc.add_paragraph(q["stem"])

            # Choices
            for c in q["choices"]:
                doc.add_paragraph(f"    {c['id']}.  {c['text']}")

            doc.add_paragraph()  # spacer

            # Build answer key entry
            correct_ids = [c["id"] for c in q["choices"] if c["correct"]]
            answer_key.append(
                f"{question_num}. {', '.join(correct_ids)} — {q['answer_explanation']}"
            )

    # --- Answer Key (new page) ---
    doc.add_page_break()
    doc.add_heading("Answer Key", level=1)

    for entry in answer_key:
        p = doc.add_paragraph(entry)
        p.paragraph_format.space_after = Pt(4)

    # --- Legend ---
    doc.add_paragraph()
    legend = doc.add_paragraph()
    run = legend.add_run("★ = Essential Standard (priority for ILEARN)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(OUTPUT_PATH)
    print(f"Exported {question_num} questions → {OUTPUT_PATH}")


if __name__ == "__main__":
    export()
